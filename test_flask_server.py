

#!/usr/bin/env python3

import os
import sys
import time
import subprocess
import requests
from docx import Document
import socket

def create_test_docx():
    """Create a test DOCX file for testing"""
    doc = Document()
    doc.add_heading('Test Document', level=1)
    doc.add_paragraph('This is a paragraph in the test document.')
    doc.add_paragraph('It contains multiple paragraphs.')

    test_file = 'test_document.docx'
    doc.save(test_file)
    return test_file

def is_port_open(host, port):
    """Check if a port is open"""
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        return s.connect_ex((host, port)) == 0

def test_flask_server():
    """Test the Flask server with a sample DOCX file"""
    # Create test file
    test_file = create_test_docx()

    try:
        print("Starting Flask server...")
        server_process = subprocess.Popen(
            ["python", "flask_server.py"],
            cwd="/workspace/s2t",
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )

        # Wait for server to start (up to 10 seconds)
        max_attempts = 10
        for attempt in range(max_attempts):
            if is_port_open("localhost", 5000):
                print(f"Server started successfully on attempt {attempt + 1}")
                break
            time.sleep(1)
        else:
            # Server didn't start
            server_process.terminate()
            stdout, stderr = server_process.communicate(timeout=5)
            print("Server failed to start:")
            print("STDOUT:", stdout.decode())
            print("STDERR:", stderr.decode())
            return

        print("Testing Flask server...")

        # Test with HTTP request directly
        url = "http://localhost:5000/upload"
        files = {'file': (test_file, open(test_file, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = requests.post(url, files=files)

        if response.status_code == 200:
            result = response.json()
            print("SUCCESS!")
            print("Markdown content:")
            print(result.get('content', '')[:200] + "..." if len(result.get('content', '')) > 200 else result.get('content', ''))
            print("\nMetadata:", result.get('metadata', {}))
        else:
            print(f"ERROR: {response.status_code}")
            print("Response:", response.text)

    finally:
        # Clean up
        os.unlink(test_file)
        print("Cleaned up test file")
        if server_process.poll() is None:
            print("Stopping server...")
            server_process.terminate()
            stdout, stderr = server_process.communicate(timeout=5)
            print("Server output:")
            print(stdout.decode())
            if stderr.decode().strip():
                print("Server errors:", stderr.decode())

if __name__ == "__main__":
    test_flask_server()

