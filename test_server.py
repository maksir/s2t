


#!/usr/bin/env python3

import os
import sys
import time
import requests
from docx import Document

def create_test_docx():
    """Create a test DOCX file for testing"""
    doc = Document()
    doc.add_heading('Test Document', level=1)
    doc.add_paragraph('This is a paragraph in the test document.')
    doc.add_paragraph('It contains multiple paragraphs.')

    test_file = 'test_document.docx'
    doc.save(test_file)
    return test_file

def test_mcp_server():
    """Test the MCP server with a sample DOCX file"""
    # Create test file
    test_file = create_test_docx()

    try:
        print("Starting MCP server...")
        os.system("cd /workspace/s2t && python mcp_server.py > server.log 2>&1 &")
        time.sleep(3)  # Give the server time to start

        print("Testing MCP server...")

        # Test with HTTP request directly
        url = "http://localhost:5000/mcp"
        files = {'file': (test_file, open(test_file, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = requests.post(url, files=files)

        if response.status_code == 200:
            result = response.json()
            print("SUCCESS!")
            print("Markdown content:")
            print(result.get('content', '')[:200] + "..." if len(result.get('content', '')) > 200 else result.get('content', ''))
            print("\nMetadata:", result.get('metadata', {}))
        else:
            print(f"ERROR: {response.status_code}")
            print("Response:", response.text)

    finally:
        # Clean up
        os.unlink(test_file)
        print("Cleaned up test file")
        print("Stopping server...")
        os.system("pkill -f 'python mcp_server.py'")

if __name__ == "__main__":
    test_mcp_server()


